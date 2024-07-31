#!/usr/bin/env python
# coding: utf-8

import streamlit as st
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_community.chat_models import ChatOllama
import docx
import fitz  # PyMuPDF
from typing import List
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
from docx import Document as DocxDocument
import time
import textwrap

# Language codes for translation
LANGUAGE_CODES = {
    "English": "en",
    "German": "de",
    "French": "fr",
    "Hindi": "hi"
}

# System messages for different translation scenarios
SYSTEM_MESSAGE_TRANSLATE = """You are an expert assistant specializing in translation between German and English. Your tasks are as follows:
Language Detection:
Identify the language of the input text. Do not answer the questions if asked in text. Just translate the question text as question itself.
Translation:
If the text is in German, translate it to English.
If the text is in English, translate it to German.
Provide accurate and contextually appropriate translations. Ensure that the translated text maintains the original meaning, type, and tone.
IMPORTANT: Provide only the translated text as output.
"""

SYSTEM_MESSAGE_MULTI_LANG = """You are an expert assistant specializing in translation between various languages. Your tasks are as follows:
Translation:
If the text is in German, translate it to English.
If the text is in English, translate it to German.
If the text is in French, translate it to English.
If the text is in English, translate it to French.
If the text is in Hindi, translate it to English.
If the text is in English, translate it to Hindi.
Provide accurate and contextually appropriate translations. Ensure that the translated text maintains the original meaning, type, and tone.
IMPORTANT: Provide only the translated text as output. Do not include any additional comments or answers.
"""


# Function for translating text
def translate_text(text: str, src_lang: str, tgt_lang: str, page_type: str) -> str:
    if not text.strip():
        return "No content to translate."

    if page_type == "Translate":
        system_message = SYSTEM_MESSAGE_TRANSLATE
    else:  # Multi-Language Translator
        system_message = SYSTEM_MESSAGE_MULTI_LANG

    system_message = system_message.replace("German", src_lang).replace("English", tgt_lang)

    llm = ChatOllama(model="llama3")
    messages = [
        SystemMessage(content=system_message),
        HumanMessage(content=text)
    ]

    translated_text = llm.invoke(messages)
    return translated_text.content.strip()


# Function for detecting the language of the text
def detect_language(text: str) -> str:
    if not text.strip():
        return "No content to detect."

    system_message = """You are an expert assistant specializing in language detection. Your task is to identify the language of the input text. Provide only the detected language as output."""

    llm = ChatOllama(model="llama3")
    messages = [
        SystemMessage(content=system_message),
        HumanMessage(content=text)
    ]

    detected_language = llm.invoke(messages)
    return detected_language.content.strip()


# Functions for handling docx and pdf files
def read_docx(file) -> str:
    try:
        doc = docx.Document(file)
        full_text: List[str] = [para.text for para in doc.paragraphs]
        return '\n'.join(full_text)
    except Exception as e:
        st.error(f"Error reading DOCX: {e}")
        return ''


def read_pdf(file) -> str:
    try:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        full_text: List[str] = [page.get_text("text") for page in doc]
        return '\n'.join(full_text).strip()
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
        return 'Error reading PDF content'


# Functions to create downloadable files
def create_pdf(content: str) -> BytesIO:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    margin = 40
    max_width = width - 2 * margin
    text_object = c.beginText(margin, height - margin)
    text_object.setFont("Helvetica", 12)

    lines = content.split('\n')
    wrapped_lines = []
    for line in lines:
        wrapped_lines.extend(textwrap.wrap(line, width=100))

    for line in wrapped_lines:
        if text_object.getY() < margin:
            c.drawText(text_object)
            c.showPage()
            text_object = c.beginText(margin, height - margin)
            text_object.setFont("Helvetica", 12)
        text_object.textLine(line)

    c.drawText(text_object)
    c.save()
    buffer.seek(0)
    return buffer


def create_docx(content: str) -> BytesIO:
    doc = DocxDocument()
    doc.add_paragraph(content)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Main app setup
st.set_page_config(page_title="Language Converter", layout="wide")

# Initialize session state
if 'german_text' not in st.session_state:
    st.session_state.german_text = ""
if 'english_text' not in st.session_state:
    st.session_state.english_text = ""
if 'original_text' not in st.session_state:
    st.session_state.original_text = ""
if 'translated_text' not in st.session_state:
    st.session_state.translated_text = ""
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'translation_history' not in st.session_state:
    st.session_state.translation_history = []

# Custom CSS for styling
st.markdown(
    """
    <style>
    .main-title {
        text-align: center;
        font-size: 3em;
        margin-top: 0.5em;
        margin-bottom: 0.5em;
        color: #4B8BBE;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        font-weight: bold;
    }
    .raised-box {
        box-shadow: 0 4px 8px 0 rgba(0,0,0,0.2);
        padding: 2em;
        border-radius: 10px;
        background-color: #f9f9f9;
        margin-bottom: 1.5em;
    }
    .section-title {
        font-size: 2em;
        color: #4B8BBE;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }
    .text-area {
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
        font-size: 1em;
    }
    .center {
        display: flex;
        justify-content: center;
        align-items: center;
        margin-top: 1em;
    }
    .logo {
        width: 200px;
        height: auto;
    }
    </style>
    """, unsafe_allow_html=True)

# Add logo to the top of the sidebar
st.sidebar.image("Eng-Man.png", use_column_width=True)
st.sidebar.title("Navigation Menu")
page = st.sidebar.selectbox("Select Page", ["Home", "Translate", "Multi-Language Translator", "ChatBot"])

# Clear translation history button (not on Home page)
if page != "Home" and st.sidebar.button("Clear Translation History"):
    st.session_state.translation_history = []
    st.sidebar.success("Translation history cleared!")

# Main container
if page == "Home":
    st.markdown('<div class="main-title">Welcome to Language Translation Application</div>', unsafe_allow_html=True)
    st.markdown("""
    <div class="raised-box">
        <h2 class="section-title">Features:</h2>
        <ul>
            <li>Text Translation between German and English</li>
            <li>Document Translation (DOCX and PDF)</li>
            <li>Language Detection</li>
            <li>Translation History</li>
            <li>Multi Language Support</li>
        </ul>
    </div>
    """, unsafe_allow_html=True)

elif page == "Translate":
    with st.container():
        st.markdown('<div class="main-title">English - German Translator</div>', unsafe_allow_html=True)

        tabs = st.tabs(["Translate Text", "Translate Document"])

        with tabs[0]:
            st.markdown('<div class="section-title">Text Translation</div>', unsafe_allow_html=True)
            st.write("Enter Text to Translate:")

            col1, _ = st.columns([3, 1])

            with col1:
                st.session_state.original_text = st.text_area("Original Text", height=100, key="original_text_area",
                                                              help="Enter the text you want to translate.",
                                                              placeholder="Type or paste text here...")

            if st.button("Translate"):
                with st.spinner('Translating...'):
                    detected_language = detect_language(st.session_state.original_text)
                    if detected_language == "German":
                        st.session_state.translated_text = translate_text(st.session_state.original_text, "German",
                                                                          "English", "Translate")
                    elif detected_language == "English":
                        st.session_state.translated_text = translate_text(st.session_state.original_text, "English",
                                                                          "German", "Translate")
                    else:
                        st.session_state.translated_text = f"Detected language: {detected_language} - Translation not supported."

                st.session_state.translation_history.append({
                    "original_text": st.session_state.original_text,
                    "translated_text": st.session_state.translated_text
                })

            st.markdown('<div class="section-title">Translated Text</div>', unsafe_allow_html=True)
            st.write("Translation will appear here:")
            st.text_area("", st.session_state.translated_text, height=100, key="translated_text_area",
                         help="Translated text will be displayed here.")

        with tabs[1]:
            st.markdown('<div class="section-title">Document Translation</div>', unsafe_allow_html=True)
            st.write("Upload DOCX or PDF File to Translate:")

            uploaded_file = st.file_uploader("Choose a file", type=["docx", "pdf"])

            if uploaded_file:
                if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    original_text = read_docx(uploaded_file)
                elif uploaded_file.type == "application/pdf":
                    original_text = read_pdf(uploaded_file)
                else:
                    st.error("Unsupported file type.")
                    original_text = None

                if original_text:
                    st.text_area("Original Document Text", value=original_text, height=200,
                                 help="Original text extracted from the uploaded document.")

                    if st.button("Translate Document"):
                        with st.spinner('Translating...'):
                            detected_language = detect_language(original_text)
                            if detected_language == "German":
                                translated_text = translate_text(original_text, "German", "English", "Translate")
                            elif detected_language == "English":
                                translated_text = translate_text(original_text, "English", "German", "Translate")
                            else:
                                translated_text = f"Detected language: {detected_language} - Translation not supported."

                        st.session_state.translation_history.append({
                            "original_text": original_text,
                            "translated_text": translated_text
                        })

                        st.markdown('<div class="section-title">Translated Document Text</div>', unsafe_allow_html=True)
                        st.text_area("", translated_text, height=200, key="translated_document_text_area",
                                     help="Translated text from the uploaded document.")

                        # Provide options to download the translated document
                        pdf_buffer = create_pdf(translated_text)
                        docx_buffer = create_docx(translated_text)

                        st.download_button(label="Download Translated PDF", data=pdf_buffer, file_name="translated_document.pdf",
                                           mime="application/pdf")

                        st.download_button(label="Download Translated DOCX", data=docx_buffer, file_name="translated_document.docx",
                                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

elif page == "Multi-Language Translator":
    st.markdown('<div class="main-title">Multi-Language Translator</div>', unsafe_allow_html=True)

    tabs = st.tabs(["Translate Text", "Translate Document"])

    with tabs[0]:
        st.markdown('<div class="section-title">Text Translation</div>', unsafe_allow_html=True)
        st.write("Select Source and Target Languages:")
        src_lang = st.selectbox("Source Language", options=list(LANGUAGE_CODES.keys()))
        tgt_lang = st.selectbox("Target Language", options=list(LANGUAGE_CODES.keys()))

        st.write("Enter Text to Translate:")

        col1, _ = st.columns([3, 1])

        with col1:
            st.session_state.original_text = st.text_area("Original Text", height=100, key="multi_original_text_area",
                                                          help="Enter the text you want to translate.",
                                                          placeholder="Type or paste text here...")

        if st.button("Translate"):
            with st.spinner('Translating...'):
                st.session_state.translated_text = translate_text(st.session_state.original_text, src_lang, tgt_lang,
                                                                  "Multi-Language Translator")

            st.session_state.translation_history.append({
                "original_text": st.session_state.original_text,
                "translated_text": st.session_state.translated_text
            })

        st.markdown('<div class="section-title">Translated Text</div>', unsafe_allow_html=True)
        st.write("Translation will appear here:")
        st.text_area("", st.session_state.translated_text, height=100, key="multi_translated_text_area",
                     help="Translated text will be displayed here.")

    with tabs[1]:
        st.markdown('<div class="section-title">Document Translation</div>', unsafe_allow_html=True)
        st.write("Upload DOCX or PDF File to Translate:")

        uploaded_file = st.file_uploader("Choose a file", type=["docx", "pdf"], key="multi_file_uploader")

        if uploaded_file:
            if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                original_text = read_docx(uploaded_file)
            elif uploaded_file.type == "application/pdf":
                original_text = read_pdf(uploaded_file)
            else:
                st.error("Unsupported file type.")
                original_text = None

            if original_text:
                st.text_area("Original Document Text", value=original_text, height=200,
                             help="Original text extracted from the uploaded document.")

                if st.button("Translate Document", key="multi_translate_document_button"):
                    with st.spinner('Translating...'):
                        st.session_state.translated_text = translate_text(original_text, src_lang, tgt_lang,
                                                                          "Multi-Language Translator")

                    st.session_state.translation_history.append({
                        "original_text": original_text,
                        "translated_text": st.session_state.translated_text
                    })

                    st.markdown('<div class="section-title">Translated Document Text</div>', unsafe_allow_html=True)
                    st.text_area("", st.session_state.translated_text, height=200, key="multi_translated_document_text_area",
                                 help="Translated text from the uploaded document.")

                    # Provide options to download the translated document
                    pdf_buffer = create_pdf(st.session_state.translated_text)
                    docx_buffer = create_docx(st.session_state.translated_text)

                    st.download_button(label="Download Translated PDF", data=pdf_buffer, file_name="translated_document.pdf",
                                       mime="application/pdf")

                    st.download_button(label="Download Translated DOCX", data=docx_buffer, file_name="translated_document.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

elif page == "ChatBot":
    st.markdown('<div class="main-title">Translation ChatBot</div>', unsafe_allow_html=True)

    st.write("Ask the ChatBot to translate text between English and German:")

    user_input = st.text_input("You:", "")

    if st.button("Send"):
        if user_input.strip():
            detected_language = detect_language(user_input)
            if detected_language == "German":
                translated_text = translate_text(user_input, "German", "English", "Translate")
            elif detected_language == "English":
                translated_text = translate_text(user_input, "English", "German", "Translate")
            else:
                translated_text = f"Detected language: {detected_language} - Translation not supported."

            st.session_state.messages.append({"user": user_input, "bot": translated_text})

    for message in st.session_state.messages:
        st.write(f"**You**: {message['user']}")
        st.write(f"**Bot**: {message['bot']}")

# Display translation history (not on Home page)
if page != "Home" and st.session_state.translation_history:
    with st.expander("Translation History"):
        for record in st.session_state.translation_history:
            st.write(f"**Original**: {record['original_text']}")
            st.write(f"**Translated**: {record['translated_text']}")
            st.markdown("---")
